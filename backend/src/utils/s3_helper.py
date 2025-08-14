import boto3
import os
from werkzeug.utils import secure_filename
import io
import json
import PyPDF2
import logging

# Enhanced document processing imports
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_bytes
    TESSERACT_AVAILABLE = True
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    PDF2IMAGE_AVAILABLE = False
    logging.warning("Tesseract OCR or pdf2image not available - image processing will be limited")

try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False
    logging.warning("xlrd not available - .xls file processing will be limited")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not available - .xlsx file processing will be limited")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - Word document processing will be limited")

logger = logging.getLogger(__name__)

class S3Helper:
    def __init__(self):
        print("Initializing S3Helper with AWS credentials...")
        try:
            self.s3 = boto3.client(
                's3',
                aws_access_key_id=os.environ.get('AWS_ACCESS_KEY_ID'),
                aws_secret_access_key=os.environ.get('AWS_SECRET_ACCESS_KEY'),
                region_name=os.environ.get('AWS_REGION', 'us-east-1')
            )
            self.bucket_name = os.environ.get('AWS_S3_BUCKET') or os.environ.get('S3_BUCKET_NAME')
            print(f"Successfully initialized S3 client. Bucket name: {self.bucket_name}")
            # Test connection - don't fail on startup if permissions are limited
            try:
                self.s3.head_bucket(Bucket=self.bucket_name)
                print("Successfully connected to S3 bucket")
            except Exception as e:
                print(f"Warning: Cannot verify bucket access on startup: {str(e)}")
                print("S3Helper will still attempt operations - this may be due to limited permissions")
        except Exception as e:
            print(f"Error initializing S3 client: {str(e)}")
            raise

    def list_files(self):
        try:
            response = self.s3.list_objects_v2(Bucket=self.bucket_name)
            files = []
            if 'Contents' in response:
                for item in response['Contents']:
                    files.append({
                        'key': item['Key'],
                        'size': item['Size'],
                        'last_modified': item['LastModified'].isoformat()
                    })
            return files
        except Exception as e:
            print(f"Error listing files: {str(e)}")
            raise

    def list_objects(self, prefix: str = None):
        """List objects with optional prefix filter"""
        try:
            # Normalize prefix for root directory access
            if prefix in ["/", "", None]:
                prefix = None
                print(f"S3Helper: Listing ALL objects in bucket (root directory)")
            else:
                # Ensure prefix doesn't start with / for S3
                if prefix.startswith('/'):
                    prefix = prefix[1:]
                print(f"S3Helper: Listing objects with prefix: '{prefix}'")
            
            if prefix:
                response = self.s3.list_objects_v2(Bucket=self.bucket_name, Prefix=prefix)
            else:
                response = self.s3.list_objects_v2(Bucket=self.bucket_name)
            
            objects = []
            if 'Contents' in response:
                print(f"S3Helper: Found {len(response['Contents'])} total objects")
                for item in response['Contents']:
                    print(f"S3Helper: Processing object: {item['Key']}")
                    # Skip directory markers (keys ending with /)
                    if not item['Key'].endswith('/'):
                        objects.append(item['Key'])
                        print(f"S3Helper: Added object: {item['Key']}")
                    else:
                        print(f"S3Helper: Skipped directory marker: {item['Key']}")
            else:
                print("S3Helper: No 'Contents' found in response")
            
            print(f"S3Helper: Returning {len(objects)} objects")
            return objects
        except Exception as e:
            print(f"Error listing objects with prefix '{prefix}': {str(e)}")
            raise

    def get_file_info(self, file_key):
        """Get file metadata and info"""
        try:
            response = self.s3.head_object(Bucket=self.bucket_name, Key=file_key)
            return {
                'ContentLength': response['ContentLength'],
                'ContentType': response.get('ContentType', ''),
                'LastModified': response['LastModified'],
                'Metadata': response.get('Metadata', {})
            }
        except Exception as e:
            print(f"Error getting file info for '{file_key}': {str(e)}")
            raise

    def upload_file(self, file):
        try:
            print(f"Starting upload for file: {file.filename}")
            filename = secure_filename(file.filename)
            print(f"Secured filename: {filename}")
            
            # Test bucket access before uploading
            try:
                self.s3.head_bucket(Bucket=self.bucket_name)
                print(f"Successfully verified bucket access for: {self.bucket_name}")
            except Exception as e:
                print(f"Error accessing bucket: {str(e)}")
                return False
            
            # Attempt upload
            print(f"Uploading to bucket: {self.bucket_name}")
            self.s3.upload_fileobj(
                file,
                self.bucket_name,
                filename,
                ExtraArgs={'ACL': 'private'}
            )
            print(f"Upload completed successfully for: {filename}")
            return True
        except Exception as e:
            print(f"Error uploading file: {str(e)}")
            print(f"Error type: {type(e)}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            return False

    def upload_file_content(self, file_content, s3_key):
        """Upload file content (bytes) directly to S3 with a specific key"""
        try:
            from io import BytesIO
            
            print(f"Starting upload for S3 key: {s3_key}")
            
            # Test bucket access before uploading
            try:
                self.s3.head_bucket(Bucket=self.bucket_name)
                print(f"Successfully verified bucket access for: {self.bucket_name}")
            except Exception as e:
                print(f"Error accessing bucket: {str(e)}")
                return False
            
            # Create file-like object from content
            file_obj = BytesIO(file_content)
            
            # Attempt upload
            print(f"Uploading to bucket: {self.bucket_name} with key: {s3_key}")
            self.s3.upload_fileobj(
                file_obj,
                self.bucket_name,
                s3_key,
                ExtraArgs={'ACL': 'private'}
            )
            print(f"Upload completed successfully for: {s3_key}")
            return True
        except Exception as e:
            print(f"Error uploading file content: {str(e)}")
            print(f"Error type: {type(e)}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            return False

    def get_file_content(self, file_key):
        """Get the content of a file from S3 with enhanced document processing support"""
        try:
            print(f"Getting content for file: {file_key}")
            response = self.s3.get_object(Bucket=self.bucket_name, Key=file_key)
            
            # Get file extension for processing decisions
            file_extension = file_key.lower().split('.')[-1] if '.' in file_key else ''
            
            # Handle different file types
            if file_extension == 'pdf':
                return self._extract_pdf_text(response['Body'])
            elif file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp', 'gif']:
                return self._extract_image_text(response['Body'], file_extension)
            elif file_extension == 'xls':
                return self._extract_xls_text(response['Body'])
            elif file_extension == 'xlsx':
                return self._extract_xlsx_text(response['Body'])
            elif file_extension in ['doc', 'docx']:
                return self._extract_docx_text(response['Body'])
            else:
                # For text files and other formats, decode the content
                content = response['Body'].read().decode('utf-8')
                print(f"Successfully retrieved content for {file_key} (length: {len(content)})")
                return content
        except Exception as e:
            print(f"Error getting file content for {file_key}: {str(e)}")
            raise

    def get_file_content_detailed(self, file_key):
        """Get the content of a file from S3 with detailed response including processing info"""
        try:
            print(f"Getting content for file: {file_key}")
            response = self.s3.get_object(Bucket=self.bucket_name, Key=file_key)
            
            # Get file extension for processing decisions
            file_extension = file_key.lower().split('.')[-1] if '.' in file_key else ''
            
            # Handle different file types
            if file_extension == 'pdf':
                content = self._extract_pdf_text(response['Body'])
                return {
                    'success': True,
                    'content': content,
                    'processing_info': f'PDF processed successfully (length: {len(content)} characters)'
                }
            elif file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp', 'gif']:
                content = self._extract_image_text(response['Body'], file_extension)
                return {
                    'success': True,
                    'content': content,
                    'processing_info': f'Image processed with OCR (length: {len(content)} characters)'
                }
            elif file_extension == 'xls':
                content = self._extract_xls_text(response['Body'])
                return {
                    'success': True,
                    'content': content,
                    'processing_info': f'Excel file processed (length: {len(content)} characters)'
                }
            elif file_extension == 'xlsx':
                content = self._extract_xlsx_text(response['Body'])
                return {
                    'success': True,
                    'content': content,
                    'processing_info': f'Excel file processed (length: {len(content)} characters)'
                }
            elif file_extension in ['doc', 'docx']:
                content = self._extract_docx_text(response['Body'])
                return {
                    'success': True,
                    'content': content,
                    'processing_info': f'Word document processed (length: {len(content)} characters)'
                }
            else:
                # For text files and other formats, decode the content
                content = response['Body'].read().decode('utf-8')
                print(f"Successfully retrieved content for {file_key} (length: {len(content)})")
                return {
                    'success': True,
                    'content': content,
                    'processing_info': f'Text file processed (length: {len(content)} characters)'
                }
        except Exception as e:
            print(f"Error getting file content for {file_key}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'processing_info': f'Failed to process {file_key}'
            }

    def _extract_pdf_text(self, pdf_stream):
        """Extract text content from PDF file with enhanced error handling"""
        try:
            import PyPDF2
            
            # Read the stream into memory for PyPDF2
            pdf_content = pdf_stream.read()
            pdf_file = io.BytesIO(pdf_content)
            
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted/password protected
            if pdf_reader.is_encrypted:
                # Try to decrypt with empty password (sometimes works)
                try:
                    pdf_reader.decrypt("")
                except:
                    # If that fails, the PDF requires a password
                    raise Exception("PDF is password protected and cannot be processed. Please provide the password or process manually.")
            
            # Check if PDF has readable content
            if len(pdf_reader.pages) == 0:
                raise Exception("PDF file appears to be empty or corrupted")
            
            # Extract text from all pages
            text_content = ""
            pages_processed = 0
            
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text and page_text.strip():
                        text_content += page_text + "\n"
                        pages_processed += 1
                    
                except Exception as page_error:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {page_error}")
                    continue
            
            # Check if we extracted any meaningful content
            if not text_content.strip():
                # Try OCR as fallback for image-based PDFs
                print("PDF appears to be image-based. Attempting OCR extraction...")
                try:
                    # Create a new BytesIO stream for OCR from the original content
                    ocr_io = io.BytesIO(pdf_content)
                    ocr_text = self._extract_pdf_text_with_ocr(ocr_io)
                    if ocr_text and ocr_text.strip():
                        print(f"Successfully extracted text using OCR (length: {len(ocr_text)})")
                        return ocr_text
                    else:
                        raise Exception("PDF file contains no extractable text content. It may be image-based or corrupted.")
                except Exception as ocr_error:
                    print(f"OCR extraction failed: {ocr_error}")
                    raise Exception("PDF file contains no extractable text content. It may be image-based or corrupted.")
            
            if pages_processed == 0:
                raise Exception("No pages could be processed successfully from the PDF")
            
            print(f"Successfully extracted text from PDF: {pages_processed}/{len(pdf_reader.pages)} pages processed (length: {len(text_content)})")
            return text_content
            
        except Exception as e:
            error_message = str(e)
            
            # Provide more specific error messages
            if "password" in error_message.lower() or "encrypted" in error_message.lower():
                print(f"PDF Password Protection Error: {error_message}")
                raise Exception(f"Password protected PDF: {error_message}")
            elif "corrupted" in error_message.lower() or "damaged" in error_message.lower():
                print(f"PDF Corruption Error: {error_message}")
                raise Exception(f"Corrupted PDF file: {error_message}")
            elif "no extractable text" in error_message.lower():
                print(f"PDF Content Error: {error_message}")
                raise Exception(f"Image-based or text-less PDF: {error_message}")
            else:
                print(f"PDF Processing Error: {error_message}")
                raise Exception(f"Failed to process PDF: {error_message}")

    def _extract_image_text(self, image_stream, file_extension):
        """Extract text from images using OCR (Tesseract)"""
        try:
            if not TESSERACT_AVAILABLE:
                raise Exception("Tesseract OCR is not available. Cannot process image files.")
            
            # Read the image stream
            image_content = image_stream.read()
            image_file = io.BytesIO(image_content)
            
            # Open image with PIL
            image = Image.open(image_file)
            
            # Convert to RGB if necessary (for better OCR results)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Enhanced OCR with multiple configurations for better text extraction
            extracted_text = self._enhanced_ocr_extraction(image)
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the image. The image may not contain readable text.")
            
            logger.info(f"Successfully extracted text from image (length: {len(extracted_text)})")
            return extracted_text.strip()
            
        except Exception as e:
            error_message = str(e)
            logger.error(f"Image OCR Error: {error_message}")
            
            if "tesseract" in error_message.lower() or "not available" in error_message.lower():
                raise Exception(f"OCR service unavailable: {error_message}")
            else:
                raise Exception(f"Failed to extract text from image: {error_message}")

    def _extract_pdf_text_with_ocr(self, pdf_stream):
        """Extract text from image-based PDF using OCR (convert PDF to images first)"""
        try:
            if not TESSERACT_AVAILABLE or not PDF2IMAGE_AVAILABLE:
                raise Exception("Tesseract OCR or pdf2image is not available. Cannot process image-based PDFs.")
            
            # Read the PDF stream
            pdf_content = pdf_stream.read()
            
            # Convert PDF pages to images with higher DPI for better OCR
            # Use higher DPI (300) for better text recognition, especially for licenses
            try:
                images = convert_from_bytes(pdf_content, dpi=300, fmt='JPEG')
            except Exception:
                # Fallback to lower DPI if memory issues
                print("High DPI conversion failed, falling back to 200 DPI")
                images = convert_from_bytes(pdf_content, dpi=200, fmt='JPEG')
            
            extracted_text = ""
            pages_processed = 0
            
            for page_num, image in enumerate(images):
                try:
                    # Convert to RGB if necessary (for better OCR results)
                    if image.mode != 'RGB':
                        image = image.convert('RGB')
                    
                    # Enhanced OCR with multiple configurations for better text extraction
                    page_text = self._enhanced_ocr_extraction(image)
                    
                    if page_text and page_text.strip():
                        extracted_text += page_text + "\n"
                        pages_processed += 1
                        print(f"OCR processed page {page_num + 1}: {len(page_text)} characters")
                    
                except Exception as page_error:
                    print(f"Warning: Could not OCR page {page_num + 1}: {page_error}")
                    continue
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the PDF using OCR. The PDF may not contain readable text.")
            
            logger.info(f"Successfully extracted text from PDF using OCR: {pages_processed}/{len(images)} pages processed (length: {len(extracted_text)})")
            return extracted_text.strip()
            
        except Exception as e:
            error_message = str(e)
            logger.error(f"PDF OCR Error: {error_message}")
            
            if "tesseract" in error_message.lower() or "not available" in error_message.lower():
                raise Exception(f"OCR service unavailable: {error_message}")
            elif "pdf2image" in error_message.lower():
                raise Exception(f"PDF to image conversion failed: {error_message}")
            else:
                raise Exception(f"Failed to extract text from PDF using OCR: {error_message}")

    def _extract_xls_text(self, xls_stream):
        """Extract text from .xls files using xlrd"""
        try:
            if not XLRD_AVAILABLE:
                raise Exception("xlrd library is not available. Cannot process .xls files.")
            
            # Read the Excel file content
            xls_content = xls_stream.read()
            xls_file = io.BytesIO(xls_content)
            
            # Open the workbook
            workbook = xlrd.open_workbook(file_contents=xls_content, encoding_override='utf-8')
            
            extracted_text = []
            
            # Process each worksheet
            for sheet_index in range(workbook.nsheets):
                worksheet = workbook.sheet_by_index(sheet_index)
                sheet_name = workbook.sheet_names()[sheet_index]
                
                extracted_text.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                # Process each row
                for row_index in range(worksheet.nrows):
                    row_data = []
                    for col_index in range(worksheet.ncols):
                        cell = worksheet.cell(row_index, col_index)
                        cell_value = str(cell.value).strip()
                        if cell_value:
                            row_data.append(cell_value)
                    
                    if row_data:  # Only add non-empty rows
                        extracted_text.append(" | ".join(row_data))
            
            final_text = "\n".join(extracted_text)
            
            if not final_text.strip():
                raise Exception("No readable content found in the .xls file.")
            
            logger.info(f"Successfully extracted text from .xls file (length: {len(final_text)})")
            return final_text
            
        except Exception as e:
            error_message = str(e)
            logger.error(f"XLS Processing Error: {error_message}")
            
            if "not available" in error_message.lower():
                raise Exception(f"XLS library unavailable: {error_message}")
            elif "encrypted" in error_message.lower() or "password" in error_message.lower():
                raise Exception(f"Password protected XLS: {error_message}")
            else:
                raise Exception(f"Failed to process .xls file: {error_message}")

    def _extract_xlsx_text(self, xlsx_stream):
        """Extract text from .xlsx files using openpyxl"""
        try:
            if not OPENPYXL_AVAILABLE:
                raise Exception("openpyxl library is not available. Cannot process .xlsx files.")
            
            # Read the Excel file content
            xlsx_content = xlsx_stream.read()
            xlsx_file = io.BytesIO(xlsx_content)
            
            # Open the workbook
            workbook = openpyxl.load_workbook(xlsx_file, read_only=True, data_only=True)
            
            extracted_text = []
            
            # Process each worksheet
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                extracted_text.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                # Process each row
                for row in worksheet.iter_rows(values_only=True):
                    row_data = []
                    for cell_value in row:
                        if cell_value is not None:
                            cell_str = str(cell_value).strip()
                            if cell_str:
                                row_data.append(cell_str)
                    
                    if row_data:  # Only add non-empty rows
                        extracted_text.append(" | ".join(row_data))
            
            workbook.close()
            final_text = "\n".join(extracted_text)
            
            if not final_text.strip():
                raise Exception("No readable content found in the .xlsx file.")
            
            logger.info(f"Successfully extracted text from .xlsx file (length: {len(final_text)})")
            return final_text
            
        except Exception as e:
            error_message = str(e)
            logger.error(f"XLSX Processing Error: {error_message}")
            
            if "not available" in error_message.lower():
                raise Exception(f"XLSX library unavailable: {error_message}")
            elif "encrypted" in error_message.lower() or "password" in error_message.lower():
                raise Exception(f"Password protected XLSX: {error_message}")
            else:
                raise Exception(f"Failed to process .xlsx file: {error_message}")

    def _extract_docx_text(self, docx_stream):
        """Extract text from .docx files using python-docx"""
        try:
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library is not available. Cannot process Word documents.")
            
            # Read the Word document content
            docx_content = docx_stream.read()
            docx_file = io.BytesIO(docx_content)
            
            # Open the document
            document = Document(docx_file)
            
            extracted_text = []
            
            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    extracted_text.append(paragraph_text)
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    
                    if row_data:
                        extracted_text.append(" | ".join(row_data))
            
            final_text = "\n".join(extracted_text)
            
            if not final_text.strip():
                raise Exception("No readable content found in the Word document.")
            
            logger.info(f"Successfully extracted text from Word document (length: {len(final_text)})")
            return final_text
            
        except Exception as e:
            error_message = str(e)
            logger.error(f"DOCX Processing Error: {error_message}")
            
            if "not available" in error_message.lower():
                raise Exception(f"DOCX library unavailable: {error_message}")
            elif "encrypted" in error_message.lower() or "password" in error_message.lower():
                raise Exception(f"Password protected Word document: {error_message}")
            else:
                raise Exception(f"Failed to process Word document: {error_message}")

    def get_supported_file_types(self):
        """Get list of supported file types with current library availability"""
        supported_types = {
            'text': ['txt', 'md', 'json', 'csv', 'log', 'py', 'js', 'html', 'xml'],
            'pdf': ['pdf'],
        }
        
        if TESSERACT_AVAILABLE:
            supported_types['images'] = ['jpg', 'jpeg', 'png', 'tiff', 'bmp', 'gif']
        
        if XLRD_AVAILABLE:
            supported_types['excel_legacy'] = ['xls']
            
        if OPENPYXL_AVAILABLE:
            supported_types['excel_modern'] = ['xlsx']
            
        if DOCX_AVAILABLE:
            supported_types['word'] = ['docx', 'doc']
        
        return supported_types

    def search_files(self, search_query, file_extensions=None):
        """Search for files based on query and optionally filter by file extensions"""
        try:
            files = self.list_files()
            matching_files = []
            
            # Get supported file types for content search
            supported_types = self.get_supported_file_types()
            searchable_extensions = []
            for type_list in supported_types.values():
                searchable_extensions.extend(type_list)
            
            for file_info in files:
                file_key = file_info['key']
                
                # Filter by file extension if specified
                if file_extensions:
                    if not any(file_key.endswith(ext) for ext in file_extensions):
                        continue
                
                # Check if search query matches filename
                if search_query.lower() in file_key.lower():
                    matching_files.append(file_info)
                    continue
                
                # For supported file types, search content
                file_extension = file_key.lower().split('.')[-1] if '.' in file_key else ''
                if file_extension in searchable_extensions:
                    try:
                        content = self.get_file_content(file_key)
                        if search_query.lower() in content.lower():
                            file_info['match_type'] = 'content'
                            matching_files.append(file_info)
                    except Exception as e:
                        print(f"Error searching content of {file_key}: {str(e)}")
                        continue
            
            return matching_files
        except Exception as e:
            print(f"Error searching files: {str(e)}")
            raise

    def get_file_summary(self, file_key):
        """Get a summary of file information"""
        try:
            # Get file metadata
            response = self.s3.head_object(Bucket=self.bucket_name, Key=file_key)
            
            summary = {
                'key': file_key,
                'size': response['ContentLength'],
                'last_modified': response['LastModified'].isoformat(),
                'content_type': response.get('ContentType', 'unknown'),
                'metadata': response.get('Metadata', {})
            }
            
            # Get supported file types for content preview
            supported_types = self.get_supported_file_types()
            previewable_extensions = []
            for type_list in supported_types.values():
                previewable_extensions.extend(type_list)
            
            # Add content preview for supported file types
            file_extension = file_key.lower().split('.')[-1] if '.' in file_key else ''
            if file_extension in previewable_extensions:
                try:
                    content = self.get_file_content(file_key)
                    summary['content_preview'] = content[:500] + ('...' if len(content) > 500 else '')
                    summary['content_length'] = len(content)
                    summary['file_type_supported'] = True
                except Exception as e:
                    summary['content_preview'] = f"Error reading content: {str(e)}"
                    summary['file_type_supported'] = False
            else:
                summary['file_type_supported'] = False
                summary['content_preview'] = "File type not supported for text extraction"
            
            return summary
        except Exception as e:
            print(f"Error getting file summary: {str(e)}")
            raise

    def _enhanced_ocr_extraction(self, image):
        """
        Optimized OCR extraction with smart early termination to avoid timeouts
        """
        try:
            import numpy as np
            from PIL import ImageEnhance, ImageFilter
        except ImportError:
            # Fallback to basic OCR if numpy/PIL enhancements not available
            return pytesseract.image_to_string(image, lang='eng')
        
        best_text = ""
        max_confidence = 0
        
        # Reduced OCR configurations - only the most effective ones
        ocr_configs = [
            # Most reliable configurations first
            '--psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,:/\\()- ',
            '--psm 3 --oem 3',
            '--psm 11 --oem 3'
        ]
        
        # Reduced preprocessing - only the most effective ones
        preprocessed_images = []
        
        # Try original first
        preprocessed_images.append(("original", image))
        
        # Enhanced contrast (most effective for licenses)
        try:
            enhancer = ImageEnhance.Contrast(image)
            contrast_img = enhancer.enhance(1.5)
            preprocessed_images.append(("enhanced_contrast", contrast_img))
        except:
            pass
        
        # Upscaled image (good for small text)
        try:
            width, height = image.size
            upscaled_img = image.resize((width * 2, height * 2), Image.LANCZOS)
            preprocessed_images.append(("upscaled", upscaled_img))
        except:
            pass
        
        # Smart OCR with early termination
        attempts = 0
        max_attempts = 9  # 3 preprocessing × 3 configs = max 9 attempts
        
        for img_name, proc_img in preprocessed_images:
            for config in ocr_configs:
                attempts += 1
                try:
                    # Extract text with current configuration
                    text = pytesseract.image_to_string(proc_img, lang='eng', config=config)
                    
                    if text and text.strip():
                        # Simple confidence estimation
                        text_length = len(text.strip())
                        # Estimate confidence based on text length and word count
                        word_count = len(text.strip().split())
                        confidence_score = min(95, text_length * 0.1 + word_count * 2)
                        
                        # Keep track of best result
                        if confidence_score > max_confidence and text_length > len(best_text):
                            max_confidence = confidence_score
                            best_text = text.strip()
                        
                        print(f"OCR attempt ({img_name}, config: {config[:15]}...): {text_length} chars, confidence: {confidence_score:.1f}")
                        
                        # Early termination if we get a very good result
                        if confidence_score > 85 and text_length > 200:
                            print(f"High-quality OCR result found early, stopping after {attempts} attempts")
                            break
                    
                except Exception as config_error:
                    continue
            
            # If we found a very good result, don't try more preprocessing
            if max_confidence > 85 and len(best_text) > 200:
                break
        
        # Final fallback to simple OCR if enhanced methods failed
        if not best_text or len(best_text.strip()) < 20:
            try:
                print("Enhanced OCR yielded poor results, falling back to simple OCR")
                fallback_text = pytesseract.image_to_string(image, lang='eng')
                if len(fallback_text.strip()) > len(best_text):
                    best_text = fallback_text.strip()
            except:
                pass
        
        print(f"Enhanced OCR final result: {len(best_text)} characters extracted in {attempts} attempts")
        return best_text
